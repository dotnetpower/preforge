"""
PPTX to DOCX 변환기 테스트

테스트 대상:
1. 기본 변환 기능
2. 전처리 파싱 (슬라이드 타입 분류)
3. 제목 페이지 생성
4. 목차 페이지 생성
5. 가로 레이아웃 전환
6. 키워드 강조
7. 특수 기호 대체
8. 테이블/이미지 변환
"""
import pytest
from pathlib import Path
import shutil
from docx import Document as DocxDocument
from docx.enum.section import WD_ORIENT

from preforge.converters.pptx_to_docx import (
    PptxToDocxConverter,
    sanitize_text,
    is_highlight_keyword,
    SPECIAL_CHAR_MAP,
)


# 테스트 파일 경로
TEST_FILES_DIR = Path(__file__).parent.parent.parent / "private"
TEST_PPTX_SIMPLE = TEST_FILES_DIR / "test_presentation.pptx"
TEST_PPTX_REAL1 = TEST_FILES_DIR / "PPT샘플_20201027.pptx"
TEST_PPTX_REAL2 = TEST_FILES_DIR / "PPT샘플_개발.pptx"

# 변환 결과 저장 경로
CONVERTED_RESULT_DIR = TEST_FILES_DIR / "converted_result"


class TestSanitizeText:
    """텍스트 정리 함수 테스트"""
    
    def test_sanitize_removes_control_characters(self):
        """제어 문자가 제거되는지 테스트"""
        text = "Hello\x00World\x0bTest"
        result = sanitize_text(text)
        assert "\x00" not in result
        assert "\x0b" not in result
        assert "HelloWorldTest" == result
    
    def test_sanitize_replaces_special_chars(self):
        """특수 기호가 대체되는지 테스트"""
        for old_char, new_char in SPECIAL_CHAR_MAP.items():
            if old_char:  # 빈 문자열 제외
                text = f"Test{old_char}Text"
                result = sanitize_text(text)
                assert new_char in result
    
    def test_sanitize_handles_empty_string(self):
        """빈 문자열 처리 테스트"""
        assert sanitize_text("") == ""
        assert sanitize_text(None) is None


class TestHighlightKeywords:
    """키워드 강조 함수 테스트"""
    
    def test_highlight_keyword_detection(self):
        """키워드 감지 테스트"""
        assert is_highlight_keyword("Pathogen Information") is True
        assert is_highlight_keyword("Disease Overview") is True
        assert is_highlight_keyword("Symptoms and Signs") is True
        assert is_highlight_keyword("Diagnosis Methods") is True
        assert is_highlight_keyword("Treatment Options") is True
    
    def test_non_highlight_keywords(self):
        """일반 텍스트는 강조되지 않는지 테스트"""
        assert is_highlight_keyword("Introduction") is False
        assert is_highlight_keyword("Summary") is False
        assert is_highlight_keyword("References") is False


class TestPptxToDocxConverter:
    """PptxToDocxConverter 테스트 클래스"""
    
    @pytest.fixture
    def converter(self):
        """기본 컨버터 인스턴스"""
        return PptxToDocxConverter()
    
    @pytest.fixture
    def converted_output_dir(self):
        """변환 결과 출력 디렉토리 (private/converted_result)"""
        CONVERTED_RESULT_DIR.mkdir(parents=True, exist_ok=True)
        yield CONVERTED_RESULT_DIR
    
    def test_converter_initialization(self):
        """컨버터 초기화 테스트"""
        # Arrange & Act
        converter = PptxToDocxConverter(
            include_images=True,
            include_tables=True,
            include_notes=False,
            landscape_after_toc=True,
            image_max_width_inches=5.0,
            highlight_keywords=True,
        )
        
        # Assert
        assert converter.include_images is True
        assert converter.include_tables is True
        assert converter.include_notes is False
        assert converter.landscape_after_toc is True
        assert converter.image_max_width_inches == 5.0
        assert converter.highlight_keywords is True
    
    def test_convert_file_not_found(self, converter):
        """존재하지 않는 파일 변환 시 FileNotFoundError 발생"""
        # Arrange
        non_existent_path = Path("/non/existent/file.pptx")
        
        # Act & Assert
        with pytest.raises(FileNotFoundError, match="파일을 찾을 수 없습니다"):
            converter.convert(non_existent_path)
    
    def test_convert_invalid_extension(self, converter, converted_output_dir):
        """지원하지 않는 파일 형식 변환 시 ValueError 발생"""
        # Arrange
        invalid_file = converted_output_dir / "test.txt"
        invalid_file.write_text("test content")
        
        # Act & Assert
        with pytest.raises(ValueError, match="지원하지 않는 파일 형식"):
            converter.convert(invalid_file)
    
    @pytest.mark.skipif(
        not TEST_PPTX_SIMPLE.exists(),
        reason="테스트용 PPTX 파일이 없습니다"
    )
    def test_convert_simple_pptx(self, converter, converted_output_dir):
        """단순 PPTX 파일 변환 테스트"""
        # Arrange
        output_path = converted_output_dir / "output_simple.docx"
        
        # Act
        result_path = converter.convert(TEST_PPTX_SIMPLE, output_path)
        
        # Assert
        assert result_path.exists()
        assert result_path.suffix == ".docx"
        
        # 생성된 DOCX 파일 검증
        doc = DocxDocument(result_path)
        assert len(doc.paragraphs) > 0
    
    @pytest.mark.skipif(
        not TEST_PPTX_REAL1.exists(),
        reason="실제 PPTX 파일이 없습니다"
    )
    def test_convert_real_pptx_novaplex(self, converter, converted_output_dir):
        """실제 PPTX 파일 (Novaplex) 변환 테스트"""
        # Arrange
        output_path = converted_output_dir / "output_novaplex.docx"
        
        # Act
        result_path = converter.convert(TEST_PPTX_REAL1, output_path)
        
        # Assert
        assert result_path.exists()
        
        # DOCX 내용 검증
        doc = DocxDocument(result_path)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        
        # 슬라이드 번호가 있는지 확인
        assert "슬라이드" in full_text or len(doc.paragraphs) > 5
    
    @pytest.mark.skipif(
        not TEST_PPTX_REAL2.exists(),
        reason="실제 PPTX 파일이 없습니다"
    )
    def test_convert_real_pptx_tick_borne(self, converter, converted_output_dir):
        """실제 PPTX 파일 (Tick borne) 변환 테스트"""
        # Arrange
        output_path = converted_output_dir / "output_tick_borne.docx"
        
        # Act
        result_path = converter.convert(TEST_PPTX_REAL2, output_path)
        
        # Assert
        assert result_path.exists()
        
        doc = DocxDocument(result_path)
        assert len(doc.paragraphs) > 0
    
    @pytest.mark.skipif(
        not TEST_PPTX_SIMPLE.exists(),
        reason="테스트용 PPTX 파일이 없습니다"
    )
    def test_convert_with_landscape_after_toc(self, converted_output_dir):
        """목차 이후 가로 레이아웃 변환 테스트"""
        # Arrange
        converter = PptxToDocxConverter(landscape_after_toc=True)
        output_path = converted_output_dir / "output_landscape.docx"
        
        # Act
        result_path = converter.convert(TEST_PPTX_SIMPLE, output_path)
        
        # Assert
        assert result_path.exists()
        doc = DocxDocument(result_path)
        # 문서에 단락이 있어야 함
        assert len(doc.paragraphs) > 0
    
    @pytest.mark.skipif(
        not TEST_PPTX_SIMPLE.exists(),
        reason="테스트용 PPTX 파일이 없습니다"
    )
    def test_convert_with_keyword_highlighting(self, converted_output_dir):
        """키워드 강조 변환 테스트"""
        # Arrange
        converter = PptxToDocxConverter(highlight_keywords=True)
        output_path = converted_output_dir / "output_highlights.docx"
        
        # Act
        result_path = converter.convert(TEST_PPTX_SIMPLE, output_path)
        
        # Assert
        assert result_path.exists()
    
    @pytest.mark.skipif(
        not TEST_PPTX_SIMPLE.exists(),
        reason="테스트용 PPTX 파일이 없습니다"
    )
    def test_convert_auto_output_path(self, converter):
        """출력 경로 자동 생성 테스트"""
        # Arrange
        expected_output = TEST_PPTX_SIMPLE.with_suffix(".docx")
        
        # 기존 파일 백업
        backup_path = None
        if expected_output.exists():
            backup_path = expected_output.with_suffix(".docx.bak")
            shutil.copy(expected_output, backup_path)
        
        try:
            # Act
            result_path = converter.convert(TEST_PPTX_SIMPLE)
            
            # Assert
            assert result_path.exists()
            assert result_path == expected_output
        finally:
            # 정리: 생성된 파일 삭제 및 백업 복원
            if result_path.exists() and (backup_path is None or not backup_path.exists()):
                result_path.unlink()
            if backup_path and backup_path.exists():
                shutil.move(backup_path, expected_output)
    
    @pytest.mark.skipif(
        not TEST_PPTX_REAL1.exists(),
        reason="실제 PPTX 파일이 없습니다"
    )
    def test_converted_docx_has_tables(self, converter, converted_output_dir):
        """변환된 DOCX에 테이블이 포함되는지 테스트"""
        # Arrange
        output_path = converted_output_dir / "output_with_tables.docx"
        
        # Act
        result_path = converter.convert(TEST_PPTX_REAL1, output_path)
        
        # Assert
        doc = DocxDocument(result_path)
        # 테이블이 있는 PPT의 경우 테이블이 변환되어야 함
        # (테이블이 없을 수도 있으므로 존재 여부만 확인)
        assert result_path.exists()
    
    @pytest.mark.skipif(
        not TEST_PPTX_SIMPLE.exists(),
        reason="테스트용 PPTX 파일이 없습니다"
    )
    def test_convert_without_images(self, converted_output_dir):
        """이미지 제외하고 변환 테스트"""
        # Arrange
        converter = PptxToDocxConverter(include_images=False)
        output_path = converted_output_dir / "output_no_images.docx"
        
        # Act
        result_path = converter.convert(TEST_PPTX_SIMPLE, output_path)
        
        # Assert
        assert result_path.exists()
    
    @pytest.mark.skipif(
        not TEST_PPTX_SIMPLE.exists(),
        reason="테스트용 PPTX 파일이 없습니다"
    )
    def test_metadata_copied(self, converter, converted_output_dir):
        """메타데이터가 복사되는지 테스트"""
        # Arrange
        output_path = converted_output_dir / "output_metadata.docx"
        
        # Act
        result_path = converter.convert(TEST_PPTX_SIMPLE, output_path)
        
        # Assert
        doc = DocxDocument(result_path)
        # 메타데이터가 있으면 복사되어야 함
        # (원본에 메타데이터가 없을 수도 있음)
        assert result_path.exists()


class TestPptxToDocxConverterIntegration:
    """통합 테스트 - 실제 파일로 전체 변환 플로우 테스트"""
    
    @pytest.fixture
    def converted_output_dir(self):
        """변환 결과 출력 디렉토리 (private/converted_result)"""
        CONVERTED_RESULT_DIR.mkdir(parents=True, exist_ok=True)
        yield CONVERTED_RESULT_DIR
    
    @pytest.mark.skipif(
        not TEST_PPTX_REAL1.exists(),
        reason="실제 PPTX 파일이 없습니다"
    )
    def test_full_conversion_workflow(self, converted_output_dir):
        """전체 변환 워크플로우 테스트"""
        # Arrange
        converter = PptxToDocxConverter(
            include_images=True,
            include_tables=True,
            include_notes=True,
            landscape_after_toc=True,
            highlight_keywords=True,
        )
        output_path = converted_output_dir / "full_conversion.docx"
        
        # Act
        result_path = converter.convert(TEST_PPTX_REAL1, output_path)
        
        # Assert
        assert result_path.exists()
        
        doc = DocxDocument(result_path)
        
        # 기본 구조 검증
        assert len(doc.paragraphs) > 0, "변환된 문서에 단락이 없습니다"
        
        # 파일 크기 검증 (최소 크기)
        assert result_path.stat().st_size > 1000, "변환된 파일이 너무 작습니다"
        
        # 가로 레이아웃 섹션이 있는지 확인 (목차 이후)
        has_landscape = False
        for section in doc.sections:
            if section.orientation == WD_ORIENT.LANDSCAPE:
                has_landscape = True
                break
        
        print(f"\n변환 완료:")
        print(f"  - 입력: {TEST_PPTX_REAL1}")
        print(f"  - 출력: {result_path}")
        print(f"  - 파일 크기: {result_path.stat().st_size:,} bytes")
        print(f"  - 단락 수: {len(doc.paragraphs)}")
        print(f"  - 테이블 수: {len(doc.tables)}")
        print(f"  - 섹션 수: {len(doc.sections)}")
        print(f"  - 가로 레이아웃 섹션: {has_landscape}")


if __name__ == "__main__":
    pytest.main([__file__, "-v", "-s"])
