"""
Word document (.docx) parser
"""
from pathlib import Path
from typing import List, Optional
import docx
from docx.document import Document as DocxDocument
from docx.table import Table as DocxTable
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
import logging

from ..core.document import (
    Document,
    DocumentType,
    DocumentMetadata,
    TextContent,
    TableContent,
    ImageContent,
    CellImage,
    CellMerge,
)
from ..core.parser import BaseParser

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """Word document parser"""
    
    def __init__(self):
        super().__init__()
        self._numbering_counters = {}  # Track counters by numId
    
    def _get_paragraph_number(self, paragraph: Paragraph) -> Optional[str]:
        """Get the automatic numbering for a paragraph"""
        try:
            if paragraph._element.pPr is None:
                return None
            
            numPr = paragraph._element.pPr.numPr
            if numPr is None:
                return None
            
            # Get numbering level and ID
            ilvl_element = numPr.ilvl
            numId_element = numPr.numId
            
            if ilvl_element is None or numId_element is None:
                return None
            
            ilvl = ilvl_element.val
            numId = numId_element.val
            
            # numId of 0 means no numbering
            if numId == 0:
                return None
            
            # Track counters by numId
            counter_key = (numId, ilvl)
            if counter_key not in self._numbering_counters:
                self._numbering_counters[counter_key] = 0
            
            self._numbering_counters[counter_key] += 1
            counter = self._numbering_counters[counter_key]
            
            # Determine number format based on level
            if ilvl == 0:
                return f"{counter}."
            elif ilvl == 1:
                # Get parent level counter
                parent_key = (numId, ilvl - 1)
                parent_counter = self._numbering_counters.get(parent_key, 1)
                return f"{parent_counter}.{counter}"
            elif ilvl == 2:
                # Get grandparent level counter
                parent_key = (numId, ilvl - 1)
                grandparent_key = (numId, ilvl - 2)
                parent_counter = self._numbering_counters.get(parent_key, 1)
                grandparent_counter = self._numbering_counters.get(grandparent_key, 1)
                return f"{grandparent_counter}.{parent_counter}.{counter}"
            else:
                return f"[{counter}]"
            
        except Exception as e:
            logger.debug(f"Failed to extract numbering info: {e}")
            return None
    
    @property
    def supported_extensions(self) -> List[str]:
        return [".docx", ".doc"]
    
    @property
    def document_type(self) -> DocumentType:
        return DocumentType.DOCX
    
    def parse(self, file_path: Path) -> Document:
        """Parse Word document"""
        self.validate_file(file_path)
        
        docx_doc = docx.Document(file_path)
        
        # Extract metadata
        metadata = self._extract_metadata(docx_doc)
        
        # Extract text
        text_contents = self._extract_text(docx_doc)
        
        # Extract tables
        tables = self._extract_tables(docx_doc)
        
        # Extract images
        images = self._extract_images(docx_doc)
        
        return Document(
            file_path=file_path,
            doc_type=self.document_type,
            metadata=metadata,
            text_contents=text_contents,
            tables=tables,
            images=images,
            raw_content=docx_doc,
        )
    
    def _extract_metadata(self, doc: DocxDocument) -> DocumentMetadata:
        """Extract metadata"""
        core_props = doc.core_properties
        
        # Calculate page count (section-based estimation)
        page_count = len(doc.sections) if doc.sections else None
        
        return DocumentMetadata(
            title=core_props.title,
            author=core_props.author,
            created_at=core_props.created,
            modified_at=core_props.modified,
            subject=core_props.subject,
            keywords=core_props.keywords.split(",") if core_props.keywords else None,
            page_count=page_count,
            properties={
                "category": core_props.category,
                "comments": core_props.comments,
                "language": core_props.language,
                "section_count": len(doc.sections),
            }
        )
    
    def _extract_text(self, doc: DocxDocument) -> List[TextContent]:
        """Extract text (including sections, headers, footers, text boxes)"""
        text_contents = []
        current_page = 1
        
        # Process by section
        for section_idx, section in enumerate(doc.sections, 1):
            # Extract header
            if section.header:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        text_contents.append(
                            TextContent(
                                text=f"[Header] {para.text}",
                                level=0,
                                style="Header",
                                page_number=current_page,
                            )
                        )
            
            # Section break indicator
            if section_idx > 1:
                current_page = section_idx  # Increase page on section change
                text_contents.append(
                    TextContent(
                        text=f"--- Section {section_idx} ---",
                        level=0,
                        style="SectionBreak",
                        page_number=current_page,
                    )
                )
        
        # Extract body paragraphs (with position info)
        position = 0
        current_page = 1  # Initialize page
        
        for element in doc.element.body:
            if isinstance(element, CT_P):
                para = Paragraph(element, doc)
                
                # Page break confirmation (check even if no text)
                has_page_break = False
                try:
                    page_breaks = para._element.findall('.//' + qn('w:br'))
                    for br in page_breaks:
                        if br.get(qn('w:type')) == 'page':
                            current_page += 1
                            position += 2000  # Increase position after page break
                            has_page_break = True
                            break
                except:
                    pass
                
                if has_page_break:
                    continue
                
                # Empty paragraphs also increase position (for image-only paragraphs)
                if not para.text.strip():
                    position += 1000
                    continue
                
                # Determine heading level from style
                level = 0
                style_name = para.style.name if para.style else ""
                
                if "Heading" in style_name:
                    try:
                        level = int(style_name.split()[-1])
                    except (ValueError, IndexError):
                        level = 1
                
                # Check for Drawing objects (text boxes, shapes)
                has_drawing = False
                try:
                    if para._element.findall('.//' + qn('w:drawing')):
                        has_drawing = True
                        style_name = "Drawing"
                except:
                    pass
                
                # Check for automatic numbering
                number_prefix = self._get_paragraph_number(para)
                text = para.text
                if number_prefix:
                    text = f"{number_prefix} {text}"
                
                text_contents.append(
                    TextContent(
                        text=text,
                        level=level,
                        style=style_name,
                        page_number=current_page,
                        position=position,
                    )
                )
                position += 1000  # Gap between paragraphs
            
            elif isinstance(element, CT_Tbl):
                # Tables are processed separately, just increase position
                position += 5000
        
        # Extract footers
        for section_idx, section in enumerate(doc.sections, 1):
            if section.footer:
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        text_contents.append(
                            TextContent(
                                text=f"[Footer] {para.text}",
                                level=0,
                                style="Footer",
                                page_number=section_idx,
                            )
                        )
        
        return text_contents
    
    def _extract_tables(self, doc: DocxDocument) -> List[TableContent]:
        """Extract tables (merged cells, nested tables, cell images supported)"""
        tables = []
        current_page = 1
        
        # Traverse entire document to track page breaks
        table_page_map = {}  # {table_idx: page_number}
        table_idx = 0
        
        for element in doc.element.body:
            if isinstance(element, CT_P):
                para = Paragraph(element, doc)
                
                # Page break confirmation
                try:
                    page_breaks = para._element.findall('.//' + qn('w:br'))
                    for br in page_breaks:
                        if br.get(qn('w:type')) == 'page':
                            current_page += 1
                            break
                except:
                    pass
            
            elif isinstance(element, CT_Tbl):
                table_page_map[table_idx] = current_page
                table_idx += 1
        
        for table_idx, table in enumerate(doc.tables):
            if not table.rows:
                continue
            
            # First row is considered header
            headers = []
            header_merges = []
            for col_idx, cell in enumerate(table.rows[0].cells):
                headers.append(cell.text.strip().replace('\n', '<br>'))
                
                # Extract header colspan
                try:
                    tc = cell._element
                    tcPr = tc.find(qn('w:tcPr'))
                    if tcPr is not None:
                        gridSpan = tcPr.find(qn('w:gridSpan'))
                        if gridSpan is not None:
                            colspan = int(gridSpan.get(qn('w:val'), 1))
                            if colspan > 1:
                                header_merges.append(CellMerge(
                                    row=0,
                                    col=col_idx,
                                    colspan=colspan,
                                    rowspan=1,
                                    is_merged=False
                                ))
                except:
                    pass
            
            # Extract remaining rows as data
            rows = []
            cell_images = []
            cell_merges = header_merges.copy()
            seen_image_ids = set()  # Prevent duplicate images
            
            # Step 1: Collect all row data and colspan
            all_rows_data = []
            all_colspan_data = {}  # {(row, col): colspan}
            all_vmerge_data = {}  # {(row, col): 'restart' or 'continue'}
            
            for row_idx, row in enumerate(table.rows[1:], start=1):
                row_data = []
                for col_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip().replace('\n', '<br>')
                    row_data.append(cell_text)
                    
                    # Collect cell properties
                    try:
                        tc = cell._element
                        tcPr = tc.find(qn('w:tcPr'))
                        if tcPr is not None:
                            # gridSpan (colspan)
                            gridSpan = tcPr.find(qn('w:gridSpan'))
                            if gridSpan is not None:
                                colspan = int(gridSpan.get(qn('w:val'), 1))
                                if colspan > 1:
                                    all_colspan_data[(row_idx, col_idx)] = colspan
                            
                            # vMerge (rowspan)
                            vMerge = tcPr.find(qn('w:vMerge'))
                            if vMerge is not None:
                                val = vMerge.get(qn('w:val'))
                                if val == 'restart':
                                    all_vmerge_data[(row_idx, col_idx)] = 'restart'
                                else:
                                    all_vmerge_data[(row_idx, col_idx)] = 'continue'
                    except:
                        pass
                
                all_rows_data.append(row_data)
            
            # Step 2: Calculate rowspan from vMerge info
            vmerge_spans = {}  # {(start_row, col): rowspan}
            for col_idx in range(len(table.rows[0].cells)):
                current_start = None
                for row_idx in range(1, len(table.rows)):
                    if (row_idx, col_idx) in all_vmerge_data:
                        if all_vmerge_data[(row_idx, col_idx)] == 'restart':
                            # Previous merge ends, new one starts
                            if current_start is not None:
                                span = row_idx - current_start
                                if span > 1:
                                    vmerge_spans[(current_start, col_idx)] = span
                            current_start = row_idx
                        # 'continue' case continues
                    else:
                        # No vMerge - end previous merge
                        if current_start is not None:
                            span = row_idx - current_start
                            if span > 1:
                                vmerge_spans[(current_start, col_idx)] = span
                            current_start = None
                
                # End merge if still in progress at the end
                if current_start is not None:
                    span = len(table.rows) - current_start
                    if span > 1:
                        vmerge_spans[(current_start, col_idx)] = span
            
            # Step 3: Create CellMerge objects
            for (row, col), colspan in all_colspan_data.items():
                cell_merges.append(CellMerge(
                    row=row,
                    col=col,
                    colspan=colspan,
                    rowspan=1,
                    is_merged=False
                ))
            
            for (row, col), rowspan in vmerge_spans.items():
                # Check if this cell already has colspan
                existing = None
                for merge in cell_merges:
                    if merge.row == row and merge.col == col:
                        existing = merge
                        break
                
                if existing:
                    existing.rowspan = rowspan
                else:
                    cell_merges.append(CellMerge(
                        row=row,
                        col=col,
                        colspan=1,
                        rowspan=rowspan,
                        is_merged=False
                    ))
            
            # Mark parts of merged cells
            for (row, col), status in all_vmerge_data.items():
                if status == 'continue':
                    cell_merges.append(CellMerge(
                        row=row,
                        col=col,
                        colspan=1,
                        rowspan=1,
                        is_merged=True
                    ))
            
            # Step 4: Extract cell images
            rows = all_rows_data
            for row_idx, row in enumerate(table.rows[1:], start=1):
                for col_idx, cell in enumerate(row.cells):
                    try:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                if hasattr(run, '_element'):
                                        # Using qn() function
                                    
                                    for drawing in drawings:
                                        # Try to extract image
                                        try:
                                            blips = drawing.findall('.//' + qn('a:blip'))
                                            
                                            # Extract all unique blips (remove duplicates)
                                            for blip in blips:
                                                embed_id = blip.get(qn('r:embed'))
                                                if embed_id and embed_id not in seen_image_ids:
                                                    seen_image_ids.add(embed_id)
                                                    try:
                                                        # Find relationship through document part
                                                        image_part = doc.part.rels[embed_id].target_part
                                                        cell_images.append(
                                                            CellImage(
                                                                row=row_idx,
                                                                col=col_idx,
                                                                data=image_part.blob,
                                                                format=image_part.content_type.split('/')[-1],
                                                                width=0,
                                                                height=0,
                                                                embed_id=embed_id,
                                                            )
                                                        )
                                                    except KeyError:
                                                        pass
                                        except Exception as e:
                                            logger.debug(f"Failed to extract table cell image: {e}")
                                            continue
                    except Exception as e:
                        logger.debug(f"Error processing table cell: {e}")
                        pass
            
            tables.append(
                TableContent(
                    headers=headers,
                    rows=rows,
                    cell_images=cell_images,
                    cell_merges=cell_merges,
                    page_number=table_page_map.get(table_idx),
                )
            )
        
        return tables
    
    def _extract_images(self, doc: DocxDocument) -> List[ImageContent]:
        """Extract images (including Drawing objects, floating images)"""
        images = []
        position = 0
        current_page = 1
        
        # Access images through relationships
        image_rels = {}
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.target_ref:
                image_rels[rel_id] = rel.target_part
        
        # Traverse all Drawing objects in document body
        for element in doc.element.body:
            if isinstance(element, CT_P):
                para = Paragraph(element, doc)
                
                # Page break confirmation
                has_page_break = False
                try:
                    page_breaks = para._element.findall('.//' + qn('w:br'))
                    for br in page_breaks:
                        if br.get(qn('w:type')) == 'page':
                            current_page += 1
                            position += 2000
                            has_page_break = True
                            break
                except:
                    pass
                
                if has_page_break:
                    continue
                
                # Search for Drawings within paragraph
                for run in para.runs:
                    try:
                        if hasattr(run, '_element'):
                            # Inline images - use qn() function
                            drawings = run._element.findall('.//' + qn('w:drawing'))
                            
                            for drawing in drawings:
                                blips = drawing.findall('.//' + qn('a:blip'))
                                
                                for blip in blips:
                                    embed_id = blip.get(qn('r:embed'))
                                    if embed_id and embed_id in image_rels:
                                        try:
                                            image_part = image_rels[embed_id]
                                            
                                            # Try to extract image dimensions
                                            width = 0
                                            height = 0
                                            extents = drawing.findall('.//' + qn('wp:extent'))
                                            if extents:
                                                width = int(extents[0].get('cx', 0))
                                                height = int(extents[0].get('cy', 0))
                                            
                                            images.append(
                                                ImageContent(
                                                    data=image_part.blob,
                                                    format=image_part.content_type.split("/")[-1],
                                                    width=width,
                                                    height=height,
                                                    page_number=current_page,
                                                    position=position,
                                                )
                                            )
                                        except Exception as e:
                                            logger.warning(f"Failed to extract image: {e}")
                                            continue
                    except Exception as e:
                        logger.debug(f"Error processing Drawing: {e}")
                        pass
                
                position += 1000
            
            elif isinstance(element, CT_Tbl):
                position += 5000
        
        return images
